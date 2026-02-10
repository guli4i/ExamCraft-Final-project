# app.py
import os
import io
import json
import datetime
import random
import sqlite3
import uuid
import logging

import torch
from flask import Flask, request, jsonify, render_template, send_file
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from pypdf import PdfReader
from flask import session
from flask import session, redirect, url_for



from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm

from openai import OpenAI

# ===================== APP =====================

app = Flask(__name__)
app.secret_key = "secret-key-for-project"

logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)
def get_user_id():
    if "user_id" not in session:
        session["user_id"] = str(uuid.uuid4())
    return session["user_id"]

DB_PATH = "exam.db"

def get_db():
    return sqlite3.connect(DB_PATH)

def init_db():
    with get_db() as conn:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT
        )
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS exams (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            exam_type TEXT,
            questions TEXT,
            created_at TEXT
        )
        """)
init_db()


# ===================== FLAN (OPEN QUESTIONS) =====================

FLAN_MODEL = "google/flan-t5-small"

tokenizer_flan = AutoTokenizer.from_pretrained(FLAN_MODEL)
flan = AutoModelForSeq2SeqLM.from_pretrained(FLAN_MODEL)

DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
flan.to(DEVICE)
flan.eval()


def _dedupe_keep_order(items):
    seen = set()
    out = []
    for x in items:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out


def generate_open_questions(text: str, count: int) -> list[str]:
    """
    Generates 'count' open questions with FLAN-T5.
    Critical fix: use max_new_tokens (NOT max_length) so output isn't cut off.
    """
    text = (text or "").strip()
    if not text:
        return []

    prompt = (
        "Generate exam questions based on the following text.\n"
        "Return only questions, one per line.\n"
        "Each question must be clear and end with a question mark.\n\n"
        f"TEXT:\n{text}"
    )

    inputs = tokenizer_flan(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=384  # keep input reasonable, leave room mentally for generation
    ).to(DEVICE)

    # Generate more than needed, then filter/dedupe
    n_return = max(count * 2, count)

    with torch.no_grad():
        outputs = flan.generate(
            **inputs,
            max_new_tokens=64,         # <-- main fix
            num_return_sequences=n_return,
            do_sample=True,
            temperature=0.9,
            top_k=50,
            top_p=0.95,
            repetition_penalty=1.1,
            no_repeat_ngram_size=3
        )

    candidates = []
    for o in outputs:
        decoded = tokenizer_flan.decode(o, skip_special_tokens=True).strip()
        if not decoded:
            continue

        # FLAN sometimes returns multiple lines; split and collect questions
        for line in decoded.split("\n"):
            q = line.strip()
            if q.endswith("?") and len(q) > 12:
                candidates.append(q)

    candidates = _dedupe_keep_order(candidates)

    # If we still don't have enough, do a few small resamples
    tries = 0
    while len(candidates) < count and tries < 3:
        tries += 1
        with torch.no_grad():
            outputs = flan.generate(
                **inputs,
                max_new_tokens=64,
                num_return_sequences=count,
                do_sample=True,
                temperature=1.0,      # slightly higher to diversify
                top_k=50,
                top_p=0.95,
                repetition_penalty=1.1,
                no_repeat_ngram_size=3
            )

        for o in outputs:
            decoded = tokenizer_flan.decode(o, skip_special_tokens=True).strip()
            for line in decoded.split("\n"):
                q = line.strip()
                if q.endswith("?") and len(q) > 12:
                    candidates.append(q)

        candidates = _dedupe_keep_order(candidates)

    return candidates[:count]
                

# ===================== GPT (MCQ) =====================

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def shuffle_mcq(q: dict) -> dict:
    pairs = list(enumerate(q["options"]))
    random.shuffle(pairs)

    new_opts = []
    new_ans = []
    for new_i, (old_i, opt) in enumerate(pairs):
        new_opts.append(opt)
        if old_i in q["answer_indices"]:
            new_ans.append(new_i)

    q["options"] = new_opts
    q["answer_indices"] = new_ans
    return q


def generate_mcq(text: str, count: int) -> list[dict]:
    text = (text or "").strip()
    if not text:
        return []

    prompt = f"""
Return ONLY valid JSON:
{{
  "questions": [
    {{
      "question": "",
      "options": ["", "", "", ""],
      "answer_indices": [],
      "explanation": ""
    }}
  ]
}}

Rules:
- Exactly {count} questions
- 1 to 3 correct answers
- Correct answers must be randomly distributed
- options WITHOUT letters
- Use ONLY the text below

TEXT:
{text}
""".strip()

    r = client.responses.create(model="gpt-4.1-mini", input=prompt)
    raw = r.output_text or ""

    # Extract JSON safely-ish
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("GPT did not return valid JSON object.")

    raw_json = raw[start:end + 1]
    data = json.loads(raw_json)

    qs = data.get("questions", [])
    if not isinstance(qs, list):
        raise ValueError("JSON missing 'questions' list.")

    cleaned = []
    for q in qs:
        if not isinstance(q, dict):
            continue
        if "question" not in q or "options" not in q or "answer_indices" not in q:
            continue
        if not isinstance(q["options"], list) or len(q["options"]) != 4:
            continue
        if not isinstance(q["answer_indices"], list):
            continue
        cleaned.append(shuffle_mcq(q))

    return cleaned[:count]


# ===================== FILE TEXT =====================

def pdf_text(file_storage) -> str:
    reader = PdfReader(file_storage)
    return "".join((p.extract_text() or "") for p in reader.pages)


def docx_text(file_storage) -> str:
    d = Document(file_storage)
    return "\n".join(p.text for p in d.paragraphs)


# ===================== HISTORY =====================
def save_history(entry_type: str, questions):
    # 👉 БЕРЁМ ID ЗАЛОГИНЕННОГО ПОЛЬЗОВАТЕЛЯ
    user_id = session.get("user_id")

    # ❌ если пользователь не залогинен — НЕ сохраняем
    if not user_id:
        return

    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO exams (user_id, exam_type, questions, created_at)
            VALUES (?, ?, ?, ?)
            """,
            (
                user_id,
                entry_type,
                json.dumps(questions, ensure_ascii=False),
                datetime.datetime.now().isoformat(timespec="seconds")
            )
        )


def load_history():
    user_id = get_user_id()
    with get_db() as conn:
        cur = conn.execute(
            "SELECT id, exam_type, questions, created_at FROM exams WHERE user_id = ? ORDER BY id DESC",
            (user_id,)
        )
        rows = cur.fetchall()

    history = []
    for r in rows:
        history.append({
            "id": r[0],
            "type": r[1],
            "questions": json.loads(r[2]),
            "time": r[3]
        })
    return history




# ===================== ROUTES (PAGES) =====================

@app.route("/")
def intro():
    return render_template("intro.html")


@app.route("/generator")
def generator_page():
    if "user_id" not in session:
        return redirect(url_for("login"))
    return render_template("index.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/history_page")
def history_page():
    return render_template("history.html", history=load_history())


@app.route("/exam/<int:exam_id>")
def view_exam(exam_id: int):
    data = load_history()
    if exam_id < 0 or exam_id >= len(data):
        return "Not found", 404
    return render_template("exam.html", exam=data[exam_id])


# ===================== API =====================
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        if not username or not password:
            return render_template("register.html", error="All fields required")

        with get_db() as conn:
            try:
                cur = conn.execute(
                    "INSERT INTO users (username, password) VALUES (?, ?)",
                    (username, password)
                )
                session["user_id"] = cur.lastrowid
                return redirect(url_for("generator_page"))
            except sqlite3.IntegrityError:
                return render_template("register.html", error="Username already exists")

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        with get_db() as conn:
            cur = conn.execute(
                "SELECT id, password FROM users WHERE username = ?",
                (username,)
            )
            row = cur.fetchone()

        if not row or row[1] != password:
            return render_template("login.html", error="Invalid credentials")

        session["user_id"] = row[0]
        return redirect(url_for("generator_page"))

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("intro"))


@app.route("/upload_file", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "no file"}), 400

    f = request.files["file"]
    filename = (f.filename or "").lower()

    if filename.endswith(".pdf"):
        return jsonify({"text": pdf_text(f)})

    if filename.endswith(".docx"):
        return jsonify({"text": docx_text(f)})

    return jsonify({"error": "bad file"}), 400


@app.route("/generate_open", methods=["POST"])
def gen_open():
    d = request.get_json(force=True) or {}
    text = d.get("text", "")
    count = int(d.get("count", 3))
    #  INPUT VALIDATION
    if not text.strip():
        return jsonify({"error": "Empty input text"}), 400

    questions = generate_open_questions(text, count)
    save_history("open", questions)

    return jsonify({"questions": questions})


@app.route("/generate_mcq", methods=["POST"])
def gen_mcq():
    d = request.get_json(force=True) or {}
    text = d.get("text", "")
    count = int(d.get("count", 3))
    if not text.strip():
        return jsonify({"error": "Empty input text"}), 400
    
    questions = generate_mcq(text, count)
    save_history("mcq", questions)

    return jsonify({"questions": questions})


# ===================== EXPORT =====================

@app.route("/download_docx", methods=["POST"])
def download_docx():
    qs = (request.get_json(force=True) or {}).get("questions", [])
    d = Document()

    for i, q in enumerate(qs, 1):
        if isinstance(q, dict):
            d.add_paragraph(f"{i}. {q.get('question', '')}")
            options = q.get("options", [])
            for j, o in enumerate(options):
                d.add_paragraph(f"{chr(65 + j)}. {o}")
            ans = ", ".join(chr(65 + idx) for idx in q.get("answer_indices", []))
            d.add_paragraph(f"Correct answers: {ans}")
            d.add_paragraph(q.get("explanation", ""))
        else:
            d.add_paragraph(f"{i}. {str(q)}")

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="exam.docx")


@app.route("/download_pdf", methods=["POST"])
def download_pdf():
    qs = (request.get_json(force=True) or {}).get("questions", [])
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm
    )

    styles = getSampleStyleSheet()
    story = []

    for i, q in enumerate(qs, 1):
        if isinstance(q, dict):
            story.append(Paragraph(f"<b>{i}. {q.get('question', '')}</b>", styles["Normal"]))
            options = q.get("options", [])
            for j, o in enumerate(options):
                story.append(Paragraph(f"{chr(65 + j)}. {o}", styles["Normal"]))
            ans = ", ".join(chr(65 + idx) for idx in q.get("answer_indices", []))
            story.append(Paragraph(f"<b>Correct answers:</b> {ans}", styles["Normal"]))
            story.append(Paragraph(q.get("explanation", ""), styles["Normal"]))
        else:
            story.append(Paragraph(f"<b>{i}. {str(q)}</b>", styles["Normal"]))

        story.append(Spacer(1, 12))

    doc.build(story)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="exam.pdf")


# ===================== RUN =====================

if __name__ == "__main__":
    app.run(debug=True)
